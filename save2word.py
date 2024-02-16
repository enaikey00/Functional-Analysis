## Titolo: python x analisi funzionale
## Descrizione: 
## - connessione a un database SQL Server
## - estrazione tabelle, campi e tipi di dato
## - scrittura su file .docx (Word)
## Obiettivo: automatizzare la parte di "AS IS" del documento di Analisi Funzionale, in cui si descrivono le tabelle e i campi esistenti.
## UPGRADE: usare API ChatGPT per generare una descrizione del database (Introduzione dell'AS IS) ed eventualmente per suggerire analisi per il "TO BE"

import json

# Read data from the JSON file
with open('columns_per_table.json', 'r') as f:
    columns_per_table = json.load(f)

from docx import Document

# Create a new Word document
doc = Document()

# Loop through each table name and its corresponding column names and data types
for table_name, columns_info in columns_per_table.items():
    # Add section header with the table name
    doc.add_heading(f"Table: {table_name}", level=1)
    
    # Create a table for column names and data types
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Column Name'
    hdr_cells[1].text = 'Data Type'
    
    # Loop through the column names and data types and add them to the table
    for column_name, data_type in columns_info:
        row_cells = table.add_row().cells
        row_cells[0].text = column_name
        row_cells[1].text = data_type

# Save the Word document
doc.save('database_structure.docx')

print("Word document created successfully.")